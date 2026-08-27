import os
import shutil
import tempfile
from decimal import Decimal

from django.test import TestCase, override_settings
from django.urls import reverse
from docx import Document

from appels_offres.models import AppelOffre, CritereGrille
from historique.models import HistoriqueAction
from pipeline.generation.pv import generer_pv
from pipeline.generation.report_paths import chemin_rapport
from pipeline.generation.rapport import generer_rapport
from proces_verbaux.models import ProcesVerbal
from soumissionnaires.models import Score, Soumissionnaire
from utilisateurs.models import Utilisateur


TEST_MEDIA_ROOT = tempfile.mkdtemp()
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _docx_text(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


@override_settings(MEDIA_ROOT=TEST_MEDIA_ROOT)
class ProcesVerbalWorkflowTests(TestCase):
    @classmethod
    def tearDownClass(cls):
        super().tearDownClass()
        shutil.rmtree(TEST_MEDIA_ROOT, ignore_errors=True)

    def setUp(self):
        self.president = Utilisateur.objects.create_user(
            username="president",
            password="secret",
            role=Utilisateur.Role.PRESIDENT,
            first_name="Awa",
            last_name="President",
        )
        self.evaluateur = Utilisateur.objects.create_user(
            username="eval",
            password="secret",
            role=Utilisateur.Role.EVALUATEUR,
        )
        self.ao = AppelOffre.objects.create(
            reference="AO-PV-001",
            titre="Acquisition de materiel informatique",
            date_publication="2026-01-01",
            statut=AppelOffre.Statut.EN_COURS,
        )
        self.retenu = Soumissionnaire.objects.create(
            appel_offre=self.ao,
            nom_entreprise="Alpha SARL",
            prix_lu_publiquement=Decimal("1500000"),
            prix_corrige=Decimal("1450000"),
            rang=1,
            statut_conformite=Soumissionnaire.StatutConformite.CONFORME_TECHNIQUE,
            statut_final=Soumissionnaire.StatutFinal.RETENU,
            adresse="Tokoin",
            telephone="90000001",
            email="alpha@example.com",
            beneficiaires_effectifs="Mme Alpha",
            nationalite="Togolaise",
        )
        self.conforme = Soumissionnaire.objects.create(
            appel_offre=self.ao,
            nom_entreprise="Beta SA",
            prix_lu_publiquement=Decimal("0"),
            prix_corrige=None,
            rang=2,
            statut_conformite=Soumissionnaire.StatutConformite.CONFORME_TECHNIQUE,
            statut_final=Soumissionnaire.StatutFinal.EN_COURS,
        )
        self.ecarte = Soumissionnaire.objects.create(
            appel_offre=self.ao,
            nom_entreprise="Gamma Ltd",
            prix_lu_publiquement=Decimal("1200000"),
            prix_corrige=Decimal("1190000"),
            rang=None,
            statut_conformite=Soumissionnaire.StatutConformite.NON_CONFORME_TECHNIQUE,
            statut_final=Soumissionnaire.StatutFinal.ECARTE,
            motif_rejet="Offre technique non conforme apres validation humaine",
        )
        self.groupe = CritereGrille.objects.create(
            appel_offre=self.ao,
            libelle="Technique",
            categorie=CritereGrille.Categorie.TECHNIQUE,
            valide=True,
        )
        self.critere = CritereGrille.objects.create(
            appel_offre=self.ao,
            libelle="Experience similaire",
            categorie=CritereGrille.Categorie.TECHNIQUE,
            type_groupe=CritereGrille.TypeGroupe.NOTABLE,
            note_max=Decimal("10"),
            valide=True,
            parent=self.groupe,
        )
        Score.objects.create(
            soumissionnaire=self.retenu,
            critere=self.critere,
            note=Decimal("9"),
            justification="Experience conforme",
            valide_par=self.evaluateur,
        )

    def test_pv_generated(self):
        pv = generer_pv(self.ao, utilisateur=self.evaluateur)

        self.assertEqual(pv.statut, ProcesVerbal.Statut.BROUILLON)
        self.assertTrue(pv.fichier.name.endswith(".docx"))
        self.assertTrue(os.path.exists(pv.fichier.path))
        self.assertGreater(os.path.getsize(pv.fichier.path), 0)
        self.assertIn("Alpha SARL", _docx_text(pv.fichier.path))

    def test_pv_downloads_docx(self):
        pv = generer_pv(self.ao, utilisateur=self.evaluateur)
        self.client.force_login(self.evaluateur)

        response = self.client.get(reverse("telecharger_pv", args=[self.ao.id]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"], DOCX_CONTENT_TYPE)
        self.assertIn(os.path.basename(pv.fichier.name), response["Content-Disposition"])

    def test_anonymous_cannot_download_pv(self):
        generer_pv(self.ao, utilisateur=self.evaluateur)

        response = self.client.get(reverse("telecharger_pv", args=[self.ao.id]))

        self.assertEqual(response.status_code, 302)

    def test_pv_uses_final_winner(self):
        pv = generer_pv(self.ao, utilisateur=self.evaluateur)

        text = _docx_text(pv.fichier.path)

        self.assertIn("Alpha SARL", text)
        self.assertIn("Tokoin", text)

    def test_pv_uses_final_amount(self):
        pv = generer_pv(self.ao, utilisateur=self.evaluateur)

        text = _docx_text(pv.fichier.path)

        self.assertIn("1 450 000", text)
        self.assertIn("0", text)

    def test_non_selected_tenderers_listed(self):
        pv = generer_pv(self.ao, utilisateur=self.evaluateur)

        text = _docx_text(pv.fichier.path)

        self.assertIn("Gamma Ltd", text)

    def test_rejection_reasons_not_listed_in_synthetic_pv(self):
        pv = generer_pv(self.ao, utilisateur=self.evaluateur)

        text = _docx_text(pv.fichier.path)

        self.assertIn("Gamma Ltd", text)
        self.assertIn("Non conforme technique", text)
        self.assertNotIn("Offre technique non conforme apres validation humaine", text)

    def test_human_final_decision_overrides_ai(self):
        self.conforme.motif_rejet = "Ancienne justification IA non conforme"
        self.conforme.statut_final = Soumissionnaire.StatutFinal.EN_COURS
        self.conforme.rang = 2
        self.conforme.save()

        pv = generer_pv(self.ao, utilisateur=self.evaluateur)
        text = _docx_text(pv.fichier.path)

        self.assertIn("Beta SA", text)
        self.assertIn("0", text)
        self.assertNotIn("Ancienne justification IA non conforme", text)

    def test_non_president_cannot_sign_pv(self):
        pv = generer_pv(self.ao, utilisateur=self.evaluateur)
        self.client.force_login(self.evaluateur)

        response = self.client.post(reverse("valider_pv", args=[pv.id]))

        self.assertEqual(response.status_code, 302)
        pv.refresh_from_db()
        self.assertEqual(pv.statut, ProcesVerbal.Statut.BROUILLON)
        self.assertIsNone(pv.valide_par)

    def test_president_can_sign_pv(self):
        pv = generer_pv(self.ao, utilisateur=self.evaluateur)
        self.client.force_login(self.president)

        response = self.client.post(reverse("valider_pv", args=[pv.id]))

        self.assertEqual(response.status_code, 302)
        pv.refresh_from_db()
        self.assertEqual(pv.statut, ProcesVerbal.Statut.VALIDE)
        self.assertEqual(pv.valide_par, self.president)
        self.assertIsNotNone(pv.date_validation)
        self.assertTrue(os.path.exists(chemin_rapport(self.ao.reference)))

    def test_rapport_available_after_validation(self):
        pv = generer_pv(self.ao, utilisateur=self.evaluateur)
        self.client.force_login(self.president)

        self.client.post(reverse("valider_pv", args=[pv.id]))

        response = self.client.get(reverse("rapports"))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, self.ao.reference)

        response = self.client.get(reverse("consulter_rapport", args=[self.ao.id]))
        self.assertEqual(response.status_code, 200)

        response = self.client.get(reverse("telecharger_rapport", args=[self.ao.id]))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"], DOCX_CONTENT_TYPE)
        self.assertIn(".docx", response["Content-Disposition"])

    def test_signed_pv_not_listed_to_sign(self):
        pv = generer_pv(self.ao, utilisateur=self.evaluateur)
        pv.statut = ProcesVerbal.Statut.VALIDE
        pv.valide_par = self.president
        pv.save()
        self.client.force_login(self.president)

        response = self.client.get(reverse("signature_pv"))

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, "AO-PV-001")

    def test_generation_validation_history_kept(self):
        pv = generer_pv(self.ao, utilisateur=self.evaluateur)
        self.client.force_login(self.president)

        self.client.post(reverse("valider_pv", args=[pv.id]))

        actions = HistoriqueAction.objects.filter(appel_offre=self.ao)
        self.assertEqual(actions.count(), 3)
        self.assertTrue(actions.filter(action__contains="Génération").exists())
        self.assertTrue(actions.filter(action__contains="Validation").exists())
        self.assertTrue(actions.filter(action__contains="rapport").exists())
        self.assertTrue(actions.filter(utilisateur_nom="eval").exists())
        self.assertTrue(actions.filter(utilisateur_nom="Awa President").exists())

    def test_pv_generation_requires_retenu(self):
        self.retenu.delete()
        self.ao.refresh_from_db()
        self.client.force_login(self.evaluateur)

        response = self.client.post(reverse("generer_pv", args=[self.ao.id]))

        self.assertEqual(response.status_code, 302)
        self.assertFalse(ProcesVerbal.objects.filter(appel_offre=self.ao).exists())

    def test_rapport_requires_pv_validate_and_classification(self):
        self.client.force_login(self.president)
        with self.assertRaisesMessage(ValueError, "proces-verbal"):
            generer_rapport(self.ao, utilisateur=self.president)

    def test_rapport_generated_after_validation(self):
        pv = generer_pv(self.ao, utilisateur=self.evaluateur)
        self.client.force_login(self.president)
        self.client.post(reverse("valider_pv", args=[pv.id]))

        path = chemin_rapport(self.ao.reference)
        self.assertTrue(os.path.exists(path))
        self.assertGreater(os.path.getsize(path), 0)
        self.assertIn("RAPPORT D'EVALUATION", _docx_text(path))

    def test_rapport_generated_when_scores_are_absent(self):
        Score.objects.filter(soumissionnaire__appel_offre=self.ao).delete()
        pv = generer_pv(self.ao, utilisateur=self.evaluateur)
        self.client.force_login(self.president)

        response = self.client.post(reverse("valider_pv", args=[pv.id]))

        self.assertEqual(response.status_code, 302)
        path = chemin_rapport(self.ao.reference)
        self.assertTrue(os.path.exists(path))
        self.assertIn("RAPPORT D'EVALUATION", _docx_text(path))

    def test_anonymous_cannot_download_rapport(self):
        pv = generer_pv(self.ao, utilisateur=self.evaluateur)
        self.client.force_login(self.president)
        self.client.post(reverse("valider_pv", args=[pv.id]))
        self.client.logout()

        response = self.client.get(reverse("telecharger_rapport", args=[self.ao.id]))

        self.assertEqual(response.status_code, 302)

    def test_president_can_request_report_generation(self):
        pv = generer_pv(self.ao, utilisateur=self.evaluateur)
        self.client.force_login(self.president)
        self.client.post(reverse("valider_pv", args=[pv.id]))

        response = self.client.post(reverse("generer_rapport", args=[self.ao.id]))
        self.assertEqual(response.status_code, 302)
        self.assertTrue(os.path.exists(chemin_rapport(self.ao.reference)))
