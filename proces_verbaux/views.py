from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.http import FileResponse
from django.utils import timezone
from django.conf import settings
from django.db import transaction
import os
import logging
from docx import Document

from appels_offres.models import AppelOffre
from historique.models import HistoriqueAction
from pipeline.generation.document_context import build_document_context
from pipeline.generation.report_paths import chemin_rapport, nom_fichier_rapport
from pipeline.generation.rapport import verifier_preconditions_rapport
from pipeline.generation.pv import generer_pv
from pipeline.generation.rapport import generer_rapport
from soumissionnaires.models import Soumissionnaire

from .models import ProcesVerbal

logger = logging.getLogger(__name__)

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _rapport_base_name(ao):
    return nom_fichier_rapport(ao.reference)


def _resolve_rapport_path(ao):
    return chemin_rapport(ao.reference)


def _lire_fichier_texte(path):
    if not path or not os.path.exists(path):
        return ""
    if path.lower().endswith(".docx"):
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)
    return ""


@login_required
def generer_pv_view(request, ao_id):
    """Generate the PV."""
    ao = get_object_or_404(AppelOffre, id=ao_id)
    classement_existe = ao.soumissionnaires.filter(rang__isnull=False).exists()
    attributaire_existe = ao.soumissionnaires.filter(
        statut_final=Soumissionnaire.StatutFinal.RETENU
    ).exists()

    if request.method == "POST":
        if not classement_existe:
            messages.error(
                request,
                "Impossible de generer le PV : aucun classement n'existe pour cet appel d'offres.",
            )
            return redirect("generer_pv", ao_id=ao_id)

        if not attributaire_existe:
            messages.error(
                request,
                "Impossible de generer le PV : aucun attributaire RETENU n'est disponible.",
            )
            return redirect("generer_pv", ao_id=ao_id)

        try:
            generer_pv(ao, utilisateur=request.user)
            messages.success(request, "PV genere avec succes !")
            return redirect("generer_pv", ao_id=ao_id)
        except Exception as exc:
            logger.exception("[PV] ERREUR lors de la generation du PV : %s", ao.reference)
            messages.error(request, f"Erreur lors de la generation : {exc}")
            return redirect("generer_pv", ao_id=ao_id)

    pv = ProcesVerbal.objects.filter(appel_offre=ao).first()
    document_context = build_document_context(ao)
    return render(
        request,
        "proces_verbaux/generer_pv.html",
        {
            "appel_offre": ao,
            "pv": pv,
            "pv_genere": pv and pv.fichier and bool(pv.fichier.name),
            "classement_existe": classement_existe,
            "attributaire_existe": attributaire_existe,
            "document_context": document_context,
        },
    )


@login_required
def telecharger_pv(request, ao_id):
    """Download the PV."""
    if request.user.role not in {"secretaire", "evaluateur", "president"}:
        messages.error(request, "Acces refuse.")
        return redirect("home")

    ao = get_object_or_404(AppelOffre, id=ao_id)
    pv = get_object_or_404(ProcesVerbal, appel_offre=ao)

    if not pv.fichier or not os.path.exists(pv.fichier.path):
        messages.error(request, "Le fichier PV est introuvable.")
        return redirect("generer_pv", ao_id=ao_id)

    return FileResponse(
        open(pv.fichier.path, "rb"),
        as_attachment=True,
        filename=os.path.basename(pv.fichier.name),
        content_type=DOCX_CONTENT_TYPE,
    )


@login_required
def signature_pv(request):
    """List PVs to be signed by the president."""
    if request.user.role != "president":
        messages.error(request, "Seul le president peut signer les PV.")
        return redirect("home")

    pvs = ProcesVerbal.objects.filter(
        statut=ProcesVerbal.Statut.BROUILLON
    ).select_related("appel_offre").order_by("-date_generation")
    return render(request, "proces_verbaux/signature_pv.html", {"pvs": pvs})


@login_required
def liste_pv(request):
    pvs = ProcesVerbal.objects.select_related("appel_offre", "valide_par").order_by(
        "appel_offre__created_at",
        "appel_offre__id",
    )
    return render(request, "proces_verbaux/liste_pv.html", {"pvs": pvs})


@login_required
def detail_pv(request, pv_id):
    """View a PV."""
    pv = get_object_or_404(
        ProcesVerbal.objects.select_related("appel_offre", "valide_par"),
        id=pv_id,
    )
    if request.user.role not in {"secretaire", "evaluateur", "president"}:
        messages.error(request, "Acces refuse.")
        return redirect("home")
    contenu = _lire_fichier_texte(pv.fichier.path) if pv.fichier else ""
    document_context = build_document_context(pv.appel_offre)
    return render(
        request,
        "proces_verbaux/detail_pv.html",
        {
            "pv": pv,
            "contenu_pv": contenu,
            "document_context": document_context,
            "peut_signer": request.user.role == "president",
            "peut_valider": request.user.role == "president",
        },
    )


@login_required
def rapports(request):
    if request.user.role not in {"secretaire", "evaluateur", "president"}:
        messages.error(request, "Acces refuse.")
        return redirect("home")
    rapports_disponibles = []
    pvs_valides = ProcesVerbal.objects.filter(
        statut=ProcesVerbal.Statut.VALIDE
    ).select_related("appel_offre").order_by("appel_offre__created_at", "appel_offre__id")
    for pv in pvs_valides:
        path = _resolve_rapport_path(pv.appel_offre)
        rapports_disponibles.append(
            {
                "appel_offre": pv.appel_offre,
                "date": pv.date_validation or pv.date_generation,
                "disponible": os.path.exists(path),
                "nom": os.path.basename(path),
            }
        )
    return render(
        request,
        "proces_verbaux/rapports.html",
        {"rapports": rapports_disponibles},
    )


@login_required
def consulter_rapport(request, ao_id):
    if request.user.role not in {"secretaire", "evaluateur", "president"}:
        messages.error(request, "Acces refuse.")
        return redirect("home")
    ao = get_object_or_404(AppelOffre, id=ao_id)
    path = _resolve_rapport_path(ao)
    if not os.path.exists(path):
        messages.error(
            request,
            f"Le rapport est introuvable. Chemin attendu : {path}",
        )
        return redirect("rapports")
    return render(
        request,
        "proces_verbaux/detail_rapport.html",
        {
            "appel_offre": ao,
            "contenu_rapport": _lire_fichier_texte(path),
            "document_context": build_document_context(ao),
        },
    )


@login_required
def telecharger_rapport(request, ao_id):
    if request.user.role not in {"secretaire", "evaluateur", "president"}:
        messages.error(request, "Acces refuse.")
        return redirect("home")
    ao = get_object_or_404(AppelOffre, id=ao_id)
    path = _resolve_rapport_path(ao)
    if not os.path.exists(path):
        messages.error(
            request,
            f"Le rapport est introuvable. Chemin attendu : {path}",
        )
        return redirect("rapports")
    return FileResponse(
        open(path, "rb"),
        as_attachment=True,
        filename=os.path.basename(path),
        content_type=DOCX_CONTENT_TYPE,
    )


@login_required
def generer_rapport_view(request, ao_id):
    if request.user.role not in {"secretaire", "evaluateur", "president"}:
        messages.error(request, "Acces refuse.")
        return redirect("home")

    if request.method != "POST":
        return redirect("rapports")

    ao = get_object_or_404(AppelOffre.objects.select_related("proces_verbal"), id=ao_id)
    pv = getattr(ao, "proces_verbal", None)
    if request.user.role != "president":
        messages.error(request, "Seul le president peut generer le rapport.")
        return redirect("detail_pv", pv_id=pv.id if pv else ao_id)

    ok, raison = verifier_preconditions_rapport(ao)
    if not ok:
        messages.error(request, raison)
        return redirect("rapports")

    try:
        full_path = generer_rapport(ao, utilisateur=request.user)
        messages.success(request, f"Rapport genere avec succes : {os.path.basename(full_path)}")
    except Exception as exc:
        logger.exception("Erreur generation rapport manuelle. ao_id=%s error=%s", ao_id, exc)
        messages.error(request, f"Impossible de generer le rapport. Erreur : {exc}")
    return redirect("rapports")


@login_required
def valider_pv(request, pv_id):
    """Validate and sign a PV."""
    if request.user.role != "president":
        messages.error(request, "Seul le president peut valider un PV.")
        return redirect("detail_pv", pv_id=pv_id)

    if request.method != "POST":
        messages.error(request, "La signature du PV doit etre confirmee.")
        return redirect("detail_pv", pv_id=pv_id)

    pv = get_object_or_404(ProcesVerbal.objects.select_related("appel_offre"), id=pv_id)
    chemin_attendu = _resolve_rapport_path(pv.appel_offre)

    try:
        with transaction.atomic():
            if pv.statut == ProcesVerbal.Statut.VALIDE:
                messages.info(request, "Ce PV est deja signe.")
                return redirect("detail_pv", pv_id=pv.id)

            pv.statut = ProcesVerbal.Statut.VALIDE
            pv.valide_par = request.user
            pv.date_validation = timezone.now()
            pv.save(update_fields=["statut", "valide_par", "date_validation"])

            ao = pv.appel_offre
            ao.statut = AppelOffre.Statut.CLOTURE
            ao.save(update_fields=["statut"])

            HistoriqueAction.objects.create(
                utilisateur=request.user,
                appel_offre=ao,
                action="Validation du PV d'attribution provisoire",
                details=f"PV #{pv.id} valide et signe.",
            )

            full_path = generer_rapport(ao, utilisateur=request.user)
            if not os.path.exists(full_path):
                raise FileNotFoundError(f"Le rapport attendu n'a pas ete trouve apres generation: {full_path}")
    except Exception as exc:
        logger.exception(
            "Echec validation PV ou generation rapport. pv_id=%s expected_report=%s error=%s",
            pv_id,
            chemin_attendu,
            exc,
        )
        messages.warning(
            request,
            f"Validation ou generation du rapport en echec. Chemin attendu: {chemin_attendu}. Erreur: {exc}",
        )
        return redirect("detail_pv", pv_id=pv_id)

    messages.success(request, "PV valide et signe avec succes !")
    return redirect("signature_pv")
