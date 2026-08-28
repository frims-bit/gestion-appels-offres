"""
Extraction de texte depuis PDF et DOCX.
Simple et robuste : texte natif d'abord, OCR en secours si besoin.
"""

import os
import warnings
import logging

# On masque les warnings inutiles de transformers/tensorflow au demarrage.
warnings.filterwarnings("ignore")

logger = logging.getLogger(__name__)

# === Imports optionnels legers (on ne plante pas si un module manque) ===
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None


# docTR est volontairement importe au dernier moment. Importer doctr.models au
# chargement de ce module peut charger scipy/ML et bloquer le demarrage Django.
DocumentFile = None
ocr_predictor = None
_ocr_model = None


def _load_doctr():
    """Importe docTR uniquement quand l'OCR est reellement utilisee."""
    global DocumentFile
    global ocr_predictor

    if DocumentFile is not None and ocr_predictor is not None:
        return

    logger.info("[OCR] Import docTR...")
    from doctr.io import DocumentFile as DoctrDocumentFile
    from doctr.models import ocr_predictor as doctr_ocr_predictor

    DocumentFile = DoctrDocumentFile
    ocr_predictor = doctr_ocr_predictor


def _get_ocr_model():
    """Charge le modele OCR docTR la premiere fois qu'on l'utilise."""
    global _ocr_model

    if _ocr_model is None:
        logger.info("[OCR] Chargement du modele (premiere utilisation, patience...)")
        _load_doctr()
        _ocr_model = ocr_predictor(
            det_arch="db_mobilenet_v3_large",
            reco_arch="crnn_mobilenet_v3_small",
            pretrained=True,
        )
    return _ocr_model


def extraire_texte(chemin):
    """
    Point d'entree principal.
    Detecte l'extension et appelle la bonne fonction.
    """
    if not os.path.exists(chemin):
        return f"[Erreur] Fichier introuvable : {chemin}"

    extension = os.path.splitext(chemin)[1].lower()

    if extension == ".pdf":
        return _extraire_pdf(chemin)
    elif extension == ".docx":
        return _extraire_docx(chemin)
    else:
        return f"[Erreur] Format non supporte : {extension}"


def _extraire_pdf(chemin):
    """
    Extrait le texte d'un PDF.
    Etape 1 : on essaie de lire le texte natif (rapide et precis).
    Etape 2 : si le PDF est scanne/image, on passe par l'OCR.
    """
    if pdfplumber is None:
        return "[Erreur] pdfplumber non installe. Faites : pip install pdfplumber"

    texte = ""
    logger.info("[PDF] recu : %s", chemin)

    try:
        with pdfplumber.open(chemin) as pdf:
            pages = pdf.pages
            logger.info("[PDF] nombre_pages : %s", len(pages))
            for index, page in enumerate(pages, start=1):
                page_text = page.extract_text() or ""
                texte += page_text + "\n"
                logger.debug(
                    "[PDF] page_%s_texte_natif : %s",
                    index,
                    "oui" if page_text.strip() else "non",
                )

        if texte.strip():
            logger.info("[OCR] texte_natif_trouve : oui (%s caracteres)", len(texte))
            logger.info("[OCR] pages_analysees : %s", len(pages))
            return texte

    except Exception as e:
        logger.warning("[PDF] echec_lecture_texte_natif : %s", e)

    logger.info("[EXTRACTION] OCR en cours...")
    logger.info("[PDF] Pas de texte natif detecte, lancement de l'OCR...")
    return _extraire_par_ocr(chemin)


def _extraire_par_ocr(chemin):
    """Lit un PDF scanne avec docTR (directement, sans pdf2image)."""
    logger.info("[OCR] OCR utilise : oui")
    try:
        model = _get_ocr_model()
        doc = DocumentFile.from_pdf(chemin)
        result = model(doc)
        logger.info("[OCR] pages_analysees : %s", len(result.pages))

        texte = ""
        for page in result.pages:
            for block in page.blocks:
                for line in block.lines:
                    for word in line.words:
                        texte += word.value + " "
                    texte += "\n"
                texte += "\n"
            texte += "\n"

        if texte.strip():
            logger.info("[OCR] texte_extrait : oui (%s caracteres)", len(texte))
            return texte

        logger.warning("[OCR] echec : aucun texte detecte")
        return "[OCR] Aucun texte detecte sur ce document."

    except ImportError as e:
        logger.warning("[OCR] echec_import_doctr : %s", e)
        return (
            "[Erreur OCR] docTR ou une dependance OCR est indisponible : "
            f"{e}. Installez docTR pour lire ces fichiers : pip install python-doctr"
        )
    except Exception as e:
        logger.exception("[OCR] echec : %s", e)
        return f"[Erreur OCR] Impossible de lire le PDF : {e}"


def _extraire_docx(chemin):
    """Extrait le texte d'un DOCX (paragraphes + tableaux)."""
    if Document is None:
        return "[Erreur] python-docx non installe. Faites : pip install python-docx"

    try:
        doc = Document(chemin)
        lignes = []

        # --- Paragraphes ---
        for para in doc.paragraphs:
            if para.text.strip():
                lignes.append(para.text.strip())

        # --- Tableaux (y compris imbriques) ---
        for idx, table in enumerate(doc.tables, start=1):
            lignes.append(f"\n--- Tableau {idx} ---")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    lignes.append(" | ".join(cells))

        texte = "\n".join(lignes)
        return texte if texte.strip() else "[DOCX] Document vide."

    except Exception as e:
        return f"[Erreur DOCX] Impossible de lire le fichier : {e}"
