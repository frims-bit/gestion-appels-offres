import os

from django.conf import settings
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PV_TITLE = "PROCES-VERBAL D'EVALUATION DES OFFRES"
RAPPORT_TITLE = "RAPPORT D'EVALUATION DES OFFRES"
NON_RENSEIGNE_VALUES = {"", None, "Non renseigne", "Non renseigné"}


def clear_document_body(document):
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def _cell_text(value):
    return "" if value in NON_RENSEIGNE_VALUES else str(value)


def _title(type_document):
    return RAPPORT_TITLE if type_document == "rapport" else PV_TITLE


def _logo_path():
    for filename in ("logo-togo.png", "logo-togo.jpg", "logo-togo.jpeg"):
        path = os.path.join(settings.BASE_DIR, "static", "images", filename)
        if os.path.exists(path):
            return path
    return None


def _style_document(document):
    for section in document.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)

    for style_name in ("Heading 1", "Heading 2"):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor(23, 32, 51)


def _shade_cell(cell, fill="F4F6F9"):
    tc_pr = cell._tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_table_borders(table):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B9C2CF")


def _set_cell_font(cell, bold=False, size=8):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(size)
            run.bold = bold


def _add_preview_header(document, context, type_document):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    left, right = table.rows[0].cells

    paragraph = left.paragraphs[0]
    logo = _logo_path()
    if logo:
        paragraph.add_run().add_picture(logo, width=Inches(0.35))
        paragraph.add_run(" ")
    paragraph.add_run("Republique Togolaise").bold = True
    paragraph.add_run(f"\n{_title(type_document)}")

    right.text = (
        f"Reference : {context['appel_offre']['reference']}\n"
        f"Date : {context['appel_offre']['date_publication']}"
    )
    for paragraph in right.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph()


def _add_title_block(document, context, type_document):
    autorite = context["appel_offre"]["autorite_contractante"]
    if autorite not in NON_RENSEIGNE_VALUES:
        intro = document.add_paragraph(str(autorite))
        intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(_title(type_document))
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(23, 32, 51)

    add_table(
        document,
        ["Element", "Valeur"],
        [
            ["Reference", context["appel_offre"]["reference"]],
            ["Objet", context["appel_offre"]["objet"]],
            ["Statut", context["appel_offre"]["statut"]],
        ],
    )


def add_section(document, title):
    heading = document.add_heading(title, level=1)
    if heading.runs:
        heading.runs[0].font.name = "Arial"
        heading.runs[0].font.size = Pt(11)
        heading.runs[0].font.bold = True


def add_table(document, headers, rows, empty_text="Aucune donnee disponible"):
    table = document.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_table_borders(table)

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(header)
        _shade_cell(cell)
        _set_cell_font(cell, bold=True)

    if rows:
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = _cell_text(value)
                _set_cell_font(cells[index])
    else:
        cells = table.add_row().cells
        cells[0].text = empty_text
        _set_cell_font(cells[0])
        for cell in cells[1:]:
            cell.text = ""
    document.add_paragraph()
    return table


def _add_common_sections(document, context, type_document):
    ao = context["appel_offre"]
    add_section(document, "I. Informations sur l'appel d'offres")
    add_table(
        document,
        ["Element", "Valeur"],
        [
            ["Reference", ao["reference"]],
            ["Objet", ao["objet"]],
            ["Date de publication", ao["date_publication"]],
            ["Statut", ao["statut"]],
        ],
    )

    add_section(document, "II. Soumissionnaires")
    if type_document == "pv":
        add_table(
            document,
            ["Entreprise", "Conformite", "Resultat", "Rang"],
            [
                [row["nom"], row["statut"], row["statut_final"], row["rang"]]
                for row in context["soumissionnaires"]
            ],
        )
    else:
        add_table(
            document,
            ["Entreprise", "Date de depot", "Statut", "Rang"],
            [
                [row["nom"], row["date_depot"], row["statut"], row["rang"]]
                for row in context["soumissionnaires"]
            ],
        )


def _add_report_detail_sections(document, context):
    add_section(document, "III. Examen de la recevabilite")
    add_table(
        document,
        ["Soumissionnaire", "Critere", "Valeur", "Statut", "Justification"],
        [
            [row["soumissionnaire"], row["critere"], row["valeur"], row["statut"], row["justification"]]
            for row in context["recevabilite"]
        ],
        "Aucune donnee de recevabilite disponible",
    )

    add_section(document, "IV. Evaluation technique")
    add_table(
        document,
        ["Soumissionnaire", "Critere", "Valeur", "Statut", "Score", "Justification"],
        [
            [row["soumissionnaire"], row["critere"], row["valeur"], row["statut"], row["score"], row["justification"]]
            for row in context["technique"]
        ],
        "Aucune donnee technique disponible",
    )

    add_section(document, "V. Evaluation financiere")
    add_table(
        document,
        ["Soumissionnaire", "Prix lu", "Prix corrige", "Statut", "Rang"],
        [
            [row["soumissionnaire"], row["prix_lu"], row["prix_corrige"], row["statut"], row["rang"]]
            for row in context["financier"]
        ],
        "Aucune donnee financiere disponible",
    )


def _add_classement(document, context, title):
    add_section(document, title)
    add_table(
        document,
        ["Rang", "Entreprise", "Statut", "Prix", "Score"],
        [
            [row["rang"], row["soumissionnaire"], row["statut"], row["prix"], row["score"]]
            for row in context["classement"]
        ],
        "Aucun classement final disponible",
    )


def _add_attributaire(document, context, title):
    add_section(document, title)
    attributaire = context.get("attributaire")
    if not attributaire:
        document.add_paragraph("Aucun attributaire retenu n'est enregistre.")
        return
    add_table(
        document,
        ["Element", "Valeur"],
        [
            ["Entreprise", attributaire["nom"]],
            ["Adresse", attributaire["adresse"]],
            ["Telephone", attributaire["telephone"]],
            ["Email", attributaire["email"]],
            ["Montant", attributaire["prix"]],
        ],
    )


def _add_conclusion(document, context, title):
    add_section(document, title)
    document.add_paragraph(context["conclusion"])


def _add_footer(document):
    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    table.rows[0].cells[0].text = "Signatures"
    table.rows[0].cells[1].text = "President de la commission"


def build_preview_docx(document, context, type_document):
    _style_document(document)
    _add_preview_header(document, context, type_document)
    _add_title_block(document, context, type_document)
    _add_common_sections(document, context, type_document)

    if type_document == "rapport":
        _add_report_detail_sections(document, context)
        _add_classement(document, context, "VI. Classement final")
        _add_attributaire(document, context, "VII. Attributaire")
        _add_conclusion(document, context, "VIII. Conclusion")
    else:
        _add_classement(document, context, "III. Classement final")
        _add_attributaire(document, context, "IV. Attributaire")
        _add_conclusion(document, context, "V. Conclusion")

    _add_footer(document)
