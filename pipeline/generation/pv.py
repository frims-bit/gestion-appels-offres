import os

from django.conf import settings
from django.core.files import File
from docx import Document as WordDocument

from historique.models import HistoriqueAction
from pipeline.generation.document_context import build_document_context
from pipeline.generation.docx_preview import build_preview_docx, clear_document_body
from proces_verbaux.models import ProcesVerbal


def generer_pv(appel_offre, utilisateur=None):
    """
    Genere le PV Word avec les memes donnees et la meme structure que le Preview PV.
    """
    context = build_document_context(appel_offre)
    if not context["classement"]:
        raise ValueError("Aucun classement final disponible pour generer le PV.")
    if not context.get("attributaire"):
        raise ValueError("Aucun attributaire RETENU n'est disponible pour generer le PV.")

    template_path = os.path.join(
        settings.BASE_DIR, "templates_docx", "template_pv_reel.docx"
    )
    document = WordDocument(template_path) if os.path.exists(template_path) else WordDocument()
    clear_document_body(document)
    build_preview_docx(document, context, "pv")

    ref_clean = (appel_offre.reference or "PV").replace("/", "_").replace("\\", "_")
    temp_path = os.path.join(settings.MEDIA_ROOT, f"PV_{ref_clean}.docx")
    os.makedirs(settings.MEDIA_ROOT, exist_ok=True)
    document.save(temp_path)

    pv, _ = ProcesVerbal.objects.get_or_create(appel_offre=appel_offre)
    with open(temp_path, "rb") as file_obj:
        pv.fichier.save(f"PV_{ref_clean}.docx", File(file_obj), save=False)
    pv.statut = ProcesVerbal.Statut.BROUILLON
    pv.valide_par = None
    pv.date_validation = None
    pv.save()

    if not pv.fichier or not os.path.exists(pv.fichier.path):
        raise IOError(f"Le fichier PV attendu n'a pas ete cree: {pv.fichier.name}")
    if os.path.getsize(pv.fichier.path) <= 0:
        raise IOError(f"Le fichier PV cree est vide: {pv.fichier.path}")
    WordDocument(pv.fichier.path)

    if os.path.exists(temp_path):
        os.remove(temp_path)

    HistoriqueAction.objects.create(
        utilisateur=utilisateur if getattr(utilisateur, "is_authenticated", False) else None,
        appel_offre=appel_offre,
        action="Génération du PV d'attribution provisoire",
        details=f"PV #{pv.id} genere avec la structure du Preview.",
    )

    return pv
